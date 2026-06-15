"""
生成郭展宏个人总结 Word 文档
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ===== 页面设置 =====
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.line_spacing = 1.5
paragraph_format.space_after = Pt(6)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 0:
            run.font.size = Pt(18)
        elif level == 1:
            run.font.size = Pt(15)
        elif level == 2:
            run.font.size = Pt(13)
    return h

def add_para(text, bold=False, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else None
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.bold = bold
    return p

# ==================== 标题 ====================
title = doc.add_heading('海外文物知识服务子系统——个人总结', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 个人信息表
info_table = doc.add_table(rows=4, cols=4, style='Table Grid')
info_data = [
    ['姓   名', '郭展宏', '学   号', ''],
    ['专   业', '', '班   级', ''],
    ['课程名称', '软件工程课程设计', '指导教师', ''],
    ['所属小组', '知识服务子系统组', '担任角色', '组长 / 数据查询模块负责人'],
]
for i, row_data in enumerate(info_data):
    for j, text in enumerate(row_data):
        cell = info_table.cell(i, j)
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# ==================== 正文 ====================
add_heading_styled('一、引言', level=1)
add_para(
    '本课程设计项目"海外藏中国文物知识管理与服务平台"由五个子系统组成，涵盖知识图谱构建、'
    '知识服务、智能问答、移动端应用及后台管理等方向。我所负责的"海外文物知识服务子系统"（子系统2）'
    '是平台面向用户的核心 Web 端入口，承担文物数据浏览、多维度查询、知识可视化展示以及用户个人信息管理等功能。'
    '作为本子系统的组长，我统筹了从需求分析、技术选型、任务分工、进度管理到最终集成联调的全过程，'
    '同时直接负责数据查询模块的设计与开发，并在项目后期主导了跨子系统联调工作。'
    '本总结将从组长管理职责、核心技术开发、跨系统协作三个方面进行回顾。'
)

# ==================== 二、组长管理职责 ====================
add_heading_styled('二、组长管理与协调工作', level=1)

add_heading_styled('2.1 项目进度管理', level=2)
add_para(
    '项目自2026年4月23日（第8周）启动，至6月11日（第15周）完成跨子系统联调，历时8周。'
    '我主持并记录了全部8次周例会，每次会议形成标准化的会议纪要文档（包含会议议程、上周计划复盘、'
    '本周工作进展、遇到的问题与风险、技术决策、下周工作计划及专题讨论共七个板块），'
    '确保团队每位成员对项目整体进度有清晰的认知。'
)
add_para(
    '在进度把控方面，我根据各阶段实际情况动态调整任务优先级：第8-9周集中完成需求分析与基础架构搭建，'
    '第10-11周推进四大模块核心功能开发，第12周转入文档撰写与质量收敛，'
    '第13周应对知识图谱子系统真实数据接入带来的兼容性挑战，第14周完成小组汇报答辩，'
    '第15周启动与后台管理子系统及知识问答子系统的联调工作。'
    '期间多次调整任务分配以应对突发情况——例如在第12周将Mock数据扩充任务从单人承担重新分配为双人并行，'
    '在第13周知识图谱子系统数据未按时交付时果断决策自行编写翻译脚本进行英文数据中文化处理。'
)

add_heading_styled('2.2 团队分工与协作', level=2)
add_para(
    '项目启动之初，我根据系统四大功能模块（数据浏览、数据查询、数据可视化、用户个人信息管理）的'
    '划分思路，结合每位成员的技术背景与兴趣，制定了"五人分做四模块加一人负责测试整合"的分工方案：'
    '刘思益负责数据浏览模块（文物列表/详情/对比）、我本人负责数据查询模块（全文搜索/高级查询/知识图谱查询）、'
    '田亮负责数据可视化模块（统计分析/知识图谱关系图/时间轴/地理分布图）、米嘉鹏负责用户个人信息管理模块'
    '（注册登录/个人资料/收藏/评论）、赵子皓负责测试整合与部署。该分工方案经全员讨论一致通过，'
    '后续未出现因分工不清导致的推诿或冲突。'
)
add_para(
    '在协作机制方面，我主导建立了以下规范：（1）采用GitHub Flow分支策略，各成员在feature分支开发，'
    '通过Pull Request合并至develop分支，由我统一进行代码审核与合并；（2）编写AGENTS.md开发规范文档，'
    '对TypeScript类型约束、函数式组件规范、TailwindCSS样式约束、命名约定等做了详细规定，'
    '从源头保证了五人并行开发时的代码风格一致性；（3）API接口文档先行——我在第9周即向知识图谱组和后台组'
    '发出《API接口需求初稿》，确保各子系统数据格式在正式编码前达成共识。'
)

add_heading_styled('2.3 文档体系构建', level=2)
add_para(
    '我全程负责了本子系统的文档体系建设工作，具体包括：（1）编写AGENTS.md开发规范（涵盖命名规范、'
    '组件规范、类型规范、Git工作流等章节），作为团队统一的编码标准；'
    '（2）撰写全部8份周会议纪要（第8周至第15周），累计约两万余字，完整记录了项目进展脉络；'
    '（3）主导三份核心文档（项目管理计划、需求规格说明书、设计报告）中数据查询模块和总体架构部分的撰写，'
    '并负责三份文档的最终汇总与格式统一；（4）编写用户使用手册和测试报告；'
    '（5）制作答辩PPT（32页）并录制完整功能演示视频；（6）更新FEATURES-COMPLETE.md功能清单，'
    '标注已完成功能42项及已知改进项6项。'
)

# ==================== 三、核心技术开发 ====================
add_heading_styled('三、核心技术开发工作', level=1)

add_heading_styled('3.1 数据查询模块——全文搜索与高级查询', level=2)
add_para(
    '数据查询模块是本子系统的核心功能之一，我独立完成了该模块的全部前端开发工作。'
    '该模块包含三个子功能页面：'
)
add_para(
    '（1）简单搜索页面（SearchPage.tsx）：实现了基于全文关键词匹配的文物检索功能。'
    '用户输入搜索词后，系统在文物名称、描述、朝代、类型、材质、博物馆等字段中进行模糊匹配，'
    '搜索结果以卡片视图展示，支持按相关度或年代排序。页面还集成了搜索历史记录（存储在localStorage中，'
    '按用户ID分区）、搜索结果CSV/JSON格式导出（使用Blob API实现客户端文件生成与下载）等功能。'
    '搜索输入框支持防抖处理（300ms延迟），避免频繁触发检索逻辑。'
)
add_para(
    '（2）高级搜索页面（AdvancedSearchPage.tsx）：提供了多字段组合查询能力，'
    '用户可按文物名称、类型（青铜器/瓷器/书画/玉器等20余类）、材质（青铜/陶瓷/丝绸/玉石等）、'
    '年代/朝代（商朝至清朝共18个朝代）、所属博物馆、地理区域等维度进行精确筛选。'
    '筛选条件以表单形式组织，各下拉选项从Mock数据中动态提取真实值而非硬编码，'
    '保证了数据变更时选项自动同步。高级搜索结果同样支持导出功能，且查询参数通过Zustand Store全局共享，'
    '用户可在浏览页与搜索页之间无缝切换而不丢失筛选状态。'
)
add_para(
    '（3）知识图谱查询页面（KnowledgeGraph.tsx）：支持两种查询模式——自然语言输入（如"唐代的所有瓷器"）'
    '和结构化表单查询（实体类型+名称+关系）。查询结果以ECharts关系图布局展示文物实体间的三元组关系'
    '（主体-谓词-客体），支持节点点击查看详情、图谱缩放与拖拽交互。页面底部以列表形式展示查询匹配的文物卡片，'
    '用户可一键跳转至文物详情页或加入对比列表。此外还实现了KG查询结果的CSV/JSON导出功能。'
)
add_para(
    '查询模块的各页面通过React Router统一管理路由（/search、/advanced-search、/knowledge-graph），'
    '我在导航栏Header组件中集成了统一的搜索入口，点击搜索图标即可跳转至搜索页面。'
    '查询模块的Mock API处理器（handlers.ts中的全文搜索、高级搜索、知识图谱查询三个方法）也由我编写，'
    '确保在真实后端未就绪时前端可独立开发与演示。'
)

add_heading_styled('3.2 筛选与排序功能', level=2)
add_para(
    '在浏览页面（BrowsePage）中，我实现了多维筛选面板（FilterPanel.tsx）和排序控件（SortControl.tsx）。'
    '筛选面板位于页面左侧，支持按文物类型、材质、年代、博物馆、地区五个维度组合筛选，'
    '每个维度以可折叠手风琴面板展示，选中条件以标签形式显示在顶部并可单独移除。'
    '排序控件支持按名称、年代、相关度三种方式排序，且排序与筛选条件可以任意组合使用。'
    '筛选和排序的状态统一存储在Zustand的artifactStore中，保证了用户在浏览页与详情页之间导航后筛选状态不丢失。'
    '在移动端适配方面，筛选面板在窄屏设备上自动折叠为底部弹出面板（Sheet组件），点击筛选按钮后从底部滑入。'
)

add_heading_styled('3.3 数据接入与兼容性修复', level=2)
add_para(
    '第13周是项目面临最大技术挑战的一周。知识图谱子系统（子系统1）提供的971条芝加哥艺术博物馆中国文物数据'
    '（chicago_museum.csv）全部为英文——文物名称、朝代、类型、材质、描述均无中文字段，'
    '且数据格式与前端此前基于Mock数据开发的假设存在显著差异。面对这一突发情况，'
    '我主导完成了以下紧急处理工作：'
)
add_para(
    '（1）编写Python翻译脚本（scripts/translate_artifacts.py）：该脚本内置20种文物类型中文映射'
    '（如"Ceramics"→"陶瓷器"、"Bronze"→"青铜器"）、80余种材质关键词映射'
    '（如"underglaze blue"→"青花"、"cloisonné"→"景泰蓝"、"silk"→"丝绸"）、'
    '18个朝代中英文映射（商朝Shang Dynasty→清朝Qing Dynasty），以及博物馆名称和地区的翻译规则。'
    '脚本对971条CSV记录逐条处理，生成中文化后的artifacts.ts文件（name→中文名，nameEn→英文原名），'
    '翻译覆盖率从初始的62%逐步优化至78%（601/971条名称成功翻译）。'
)
add_para(
    '（2）图片处理方案设计与实现：由于CSV数据中image_path列全部为空（0/971），图片仅提供IIIF远程URL且无本地缓存。'
    '我编写了imageUtils.ts工具模块，统一处理文物图片URL的解析、IIIF尺寸参数调整、缩略图生成及多级降级逻辑。'
    '同时实现了AuthImage组件——对需要认证的API图片使用fetch+Authorization header获取并转为blob URL渲染，'
    '对公开的IIIF/CDN图片直接作为img src加载。图片加载失败时自动尝试备用URL，全部失败则显示Gem图标占位。'
    '针对答辩现场网络不确定性，我提前下载了20张关键文物图片作为本地缓存备用。'
)
add_para(
    '（3）四大可视化模块兼容性修复：数据从英文切换为中文后，原有的统计分析、知识图谱、时间轴、地理分布四个页面'
    '出现了朝代识别失败（正则表达式仅匹配英文Dynasty模式）、博物馆坐标缺失（仅含英文博物馆名）、'
    '图表数据为空等问题。我逐一修复了以下内容：extractDynastyName函数增加中文朝代正则'
    '（/([^\\s]+朝)/）；visualizationService.ts全面重写为从真实文物数据动态计算统计分布；'
    '时间轴页面添加了18个中文朝代到公元年份的映射表；地理分布页面添加了10个博物馆的中文名称与经纬度坐标映射。'
)
add_para(
    '（4）Mock处理器及组件增强：handlers.ts中的时代提取逻辑增加了对中文朝代正则的支持；'
    'ArtifactCard、ArtifactListItem、DetailPage、ComparePage四个图片渲染组件统一增加了'
    '备用图片重试机制，并统一使用imageUtils工具模块处理图片URL。'
)

add_heading_styled('3.4 路由与导航架构', level=2)
add_para(
    '我负责了本子系统完整的路由架构设计（router.tsx），定义了全部15条路由规则，'
    '包含首页、文物浏览、文物详情、文物对比、简单搜索、高级搜索、知识图谱查询、统计分析、'
    '时间轴、地理分布图、个人资料、收藏管理、浏览记录、用户登录、用户注册等页面。'
    '需要登录认证的页面统一使用ProtectedRoute组件包裹，未登录用户访问时自动重定向至登录页并记录回跳地址。'
    '路由采用React Router v6的createBrowserRouter API，嵌套在PageLayout布局组件中，实现了Header和Footer的统一管理。'
)

add_heading_styled('3.5 接口服务层设计', level=2)
add_para(
    '为支撑后续与真实后端的平滑切换，我设计了Adapter模式的接口服务层：'
    'artifactService.ts封装了所有文物数据相关的API调用（列表查询、详情获取、全文搜索、高级搜索、'
    '知识图谱查询、结果导出、以图搜图等），内部通过API_BASE_URL区分开发环境（Vite代理）和生产环境（直连）。'
    'authService.ts封装了用户认证API（登录/注册），兼容后台管理子系统返回的多种JSON响应格式'
    '（嵌套data.token、扁平token、access_token等字段），通过mapAuthResponse函数统一转换为前端AuthResponse格式。'
    '所有需要认证的API请求自动从localStorage读取auth_token并附加Authorization: Bearer请求头，'
    '收到401响应时通过全局自定义事件（auth:unauthorized）触发自动登出与重定向。'
)

# ==================== 四、跨系统协作 ====================
add_heading_styled('四、跨系统协作与联调', level=1)

add_heading_styled('4.1 联调方案设计', level=2)
add_para(
    '第15周（小组汇报后），我主导启动了与后台管理子系统（子系统5）及知识问答子系统（子系统3）的联调工作。'
    '联调遵循"主动推进、Adapter兼容、不影响已有功能"的原则：在authService.ts和artifactService.ts中'
    '保留Adapter兼容层，同时支持Mock数据和真实API，确保单个子系统不可用时不影响本系统的其他功能。'
    '联调优先级按P0（用户认证）→ P1（文物数据）→ P2（知识问答入口与token传递）→ P3（数据迁移与统计看板）的顺序推进。'
)

add_heading_styled('4.2 与后台管理子系统对接', level=2)
add_para(
    '与子系统5的联调包括三个方面：（1）用户认证接口——将原有的localStorage Mock认证替换为真实后台API'
    '（POST /api/user/auth/login和POST /api/user/auth/register），统一token格式，'
    '在authService.ts的mapAuthResponse中兼容后台的access_token字段与本系统的token字段两种命名；'
    '（2）文物管理接口——对接后台的文物增删改查接口，确保本系统前端展示的文物数据与后台数据库一致；'
    '（3）用户数据同步——收藏、评论、浏览记录等数据从浏览器localStorage迁移至后台数据库，'
    '采用"首次全量导入+后续增量同步"的迁移策略，保留客户端本地副本作为回滚备份。'
)

add_heading_styled('4.3 与知识问答子系统对接', level=2)
add_para(
    '与子系统3的联调核心是跨系统认证token传递与跳转入口开发。由于本子系统与知识问答子系统部署在不同域名下，'
    'localStorage无法跨域共享，我设计了URL参数传递方案——在Header组件中新增"知识问答"跳转按钮，'
    '点击时从localStorage读取auth_token，经encodeURIComponent编码后附加到问答子系统URL的?token=参数中'
    '（如 http://10.4.163.140:5173/?token=<encoded_token>），问答子系统从URL参数读取token后存入自身localStorage，'
    '完成认证状态同步。桌面端显示为金色边框按钮（MessageCircle图标+文字），移动端在汉堡菜单中显示为独立菜单项。'
    '我还向知识问答组提出了安全建议——在读取token后使用window.history.replaceState清除URL中的token参数，'
    '避免token残留在浏览器历史记录或服务器访问日志中。'
)

add_heading_styled('4.4 接口规范对齐与问题追踪', level=2)
add_para(
    '在联调过程中，我识别并记录了多项跨系统接口规范差异问题：后台token字段命名不一致（access_token vs token）、'
    '知识问答子系统文物ID格式不匹配（entity_12345 vs artifact_001）、三个子系统API基础路径不统一'
    '（/api/qa vs /api/admin vs /api）、分页参数格式差异（page/size vs offset/limit）等。'
    '这些问题已被整理为《跨子系统接口规范对齐表》，并分发给各子系统负责人协调统一。'
    '同时通过Vite代理配置（vite.config.ts）实现了开发环境下对三个子系统API的统一转发，'
    '确保前端开发调试不受跨域限制。'
)

# ==================== 五、技术选型与架构设计 ====================
add_heading_styled('五、技术选型与架构设计贡献', level=1)
add_para(
    '在项目启动阶段（第8周），我组织团队讨论了前端技术栈选型方案，最终确定采用React 18 + TypeScript 5 + '
    'Vite 5 + TailwindCSS 3 + shadcn/ui组件库的技术组合。选型考量如下：React函数式组件+Hooks范式在团队中接受度高、'
    'TypeScript严格模式（零any约束）可有效减少运行时错误、Vite的HMR热更新提升开发效率、'
    'TailwindCSS原子化CSS避免样式冲突、shadcn/ui提供可定制的无障碍基础组件。'
    'UI设计风格参考大英博物馆官网（britishmuseum.org），采用深色主色调加金色点缀的视觉方案，'
    '由我在TailwindCSS配置中定义了博物馆风格的设计令牌（design tokens）。'
)
add_para(
    '在架构设计方面，我主导确定了四层架构方案：展示层（Pages）→ 组件层（Components）→ '
    '状态管理层（Zustand Store）→ 服务层（Services），各层职责清晰、依赖方向单一。'
    '其中Zustand被选为状态管理方案而非Redux，原因是本项目的状态结构相对简单（文物列表/筛选/用户认证/收藏），'
    'Zustand的轻量化API和内置持久化中间件更匹配需求。架构设计详情已记录在设计报告（plans/设计报告.md）'
    '的系统架构图与组件依赖关系中。'
)

# ==================== 六、经验与反思 ====================
add_heading_styled('六、经验与反思', level=1)

add_heading_styled('6.1 沟通协作方面的收获', level=2)
add_para(
    '作为组长，我深刻体会到在多人协作的软件项目中，"规范的建立比代码本身更重要"。'
    'AGENTS.md规范文档在项目初期投入的半天时间，在整个开发周期中持续产生回报——'
    '五人并行开发两个月，代码风格始终保持一致，合并冲突率远低于预期。每周例会的标准化流程'
    '（固定的七大板块议程）也让团队养成了"计划-执行-复盘"的节奏感，即使面临第13周数据接入这样的突发挑战，'
    '团队也能快速评估影响并制定应对方案，未出现方向性的混乱。'
)

add_heading_styled('6.2 技术能力的提升', level=2)
add_para(
    '在技术层面，本次项目让我在全栈协调与架构设计方面得到了显著锻炼。数据查询模块的开发加深了我对'
    'React状态管理（Zustand）、路由设计（React Router v6）、防抖优化等前端技术的理解；'
    '第13周的数据接入工作则是一次真实的"脏数据治理"实践——从CSV解析、Python翻译脚本编写、'
    '正则表达式匹配中文朝代、到图片IIIF协议适配和AuthImage组件设计，覆盖了完整的数据处理链路。'
    '跨系统联调阶段让我首次直面多子系统协作的复杂性——接口规范不一致、token格式差异、'
    '跨域存储隔离等问题都需要在技术方案与沟通协调之间找到平衡点。'
)

add_heading_styled('6.3 不足与改进方向', level=2)
add_para(
    '回顾整个项目周期，我认为存在以下可改进之处：'
    '（1）与知识图谱子系统的沟通启动偏晚——如果第11周就开始对接而非等到第13周，数据质量问题可以更早暴露并解决；'
    '（2）Mock数据扩充计划执行不力——第12周制定的数据扩充目标（50条以上）未能按时完成，'
    '部分原因是低估了数据编写的工作量，也反映出我在任务拆分时对工时估算不够准确；'
    '（3）翻译脚本的文物专业术语准确性未达专业水平——受限于领域知识，部分术语翻译依赖程序化规则匹配，'
    '缺乏人工校对环节，后续如有机会应将翻译结果交由文物专业人员审核。'
    '以上经验将在今后的工程项目中持续指导我的决策与行动。'
)

# ==================== 保存 ====================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '个人总结_郭展宏.docx')
doc.save(output_path)
print(f'文档已生成: {output_path}')

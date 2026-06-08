"""
生成汇报展示用的 Word 文档：
  - 组长部分：数据查询模块代码介绍（非技术人员也能理解）
  - 整个项目文件与架构介绍
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_cell_shading(cell, color):
    """设置表格单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_heading(doc, text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return heading


def add_body(doc, text):
    """添加正文段落"""
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    p.style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    for run in p.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def add_code_block(doc, code_text):
    """添加代码块样式段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p


def add_key_value(doc, key, value):
    """添加键值对样式"""
    p = doc.add_paragraph()
    run_key = p.add_run(f"▸ {key}：")
    run_key.bold = True
    run_key.font.size = Pt(11)
    run_key.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    run_val = p.add_run(value)
    run_val.font.size = Pt(11)
    run_val.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def add_table_with_style(doc, headers, rows, col_widths=None):
    """添加美观表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1a1a1a')

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
            if r % 2 == 1:
                set_cell_shading(cell, 'f9f9f4')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # 表后空行
    return table


def build_document():
    doc = Document()

    # ─── 全局字体设置 ───
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ======================= 封面 =======================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('海外藏中国文物知识管理与服务平台')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0xC9, 0xA9, 0x61)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('数据查询模块 · 汇报展示文档')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'文档生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}\n'
                        '角色：项目组长\n'
                        '负责模块：(2) 数据查询')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_page_break()

    # ======================= 目录页 =======================
    add_styled_heading(doc, '目  录', level=1)
    toc_items = [
        '第一部分：我的模块 —— 数据查询功能详解',
        '  1.1 模块概述',
        '  1.2 简单查询（全文检索）',
        '  1.3 高级查询（组合条件筛选）',
        '  1.4 知识图谱查询（关系探索 + 自然语言）',
        '  1.5 查询结果导出（CSV / JSON）',
        '  1.6 我的模块涉及的核心文件一览',
        '',
        '第二部分：整个项目的文件架构介绍',
        '  2.1 项目技术栈总览',
        '  2.2 项目文件结构（逐文件说明）',
        '  2.3 页面路由一览',
        '  2.4 数据如何流转（通俗版）',
        '  2.5 功能模块总览',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ======================= 第一部分 =======================
    add_styled_heading(doc, '第一部分：我的模块 —— 数据查询功能详解', level=1)

    # --- 1.1 ---
    add_styled_heading(doc, '1.1  模块概述', level=2)

    add_body(doc,
        '本模块是整个"海外藏中国文物知识管理与服务平台"的 核心入口模块之一——数据查询。\n\n'
        '打个比方：如果把平台比作一个大型数字博物馆，数据查询模块就是博物馆的"问询台 + 检索终端"。\n'
        '用户来到平台，最自然的需求就是"找到我想看的文物"。本模块提供了从简单到复杂的多层次查询能力，'
        '让不同需求的用户都能快速定位到目标文物。'
    )

    add_body(doc, '模块包含四个子功能：')
    add_key_value(doc, '① 简单查询', '像百度/谷歌一样，输入关键词就能搜索文物（支持按名称、博物馆、年代搜索）')
    add_key_value(doc, '② 高级查询', '像电商筛选商品一样，通过多个下拉框组合限定查询条件（类型、年代范围、收藏地等）')
    add_key_value(doc, '③ 知识图谱查询', '以图形化的方式展示文物之间的关联关系，也支持用自然语言提问（如"唐朝的瓷器"）')
    add_key_value(doc, '④ 查询结果导出', '将搜索结果下载为 CSV 表格或 JSON 数据文件，方便进一步研究')

    add_body(doc,
        '下面用"非技术人员也能听懂"的方式，逐一介绍每个子功能的实现思路。'
    )

    # --- 1.2 ---
    add_styled_heading(doc, '1.2  简单查询（全文检索）', level=2)

    add_body(doc,
        '📄 对应文件：src/pages/SearchPage.tsx（约 300 行代码）\n\n'
        '【功能通俗描述】\n'
        '用户在搜索框中输入关键词（比如"石碑"、"大英博物馆"、"唐代"），点击搜索后，系统会从数据库中'
        '检索所有匹配的文物记录，并以卡片形式展示出来。\n\n'
        '【核心功能点】'
    )

    points = [
        ('搜索框 + 快捷搜索', '页面顶部提供一个大型搜索框，用户输入关键词后回车或点击"搜索"按钮即可发起查询。'
         '同时还显示最近的搜索历史，点击历史记录可以快速重复搜索。'),
        ('搜索结果展示', '查询结果以卡片网格的形式排列（每行 4 张卡片），每张卡片显示文物的图片、'
         '名称、年代、地区、材质等关键信息，方便用户快速浏览。'),
        ('分页浏览', '结果数量较多时自动分页（每页 20 条），底部提供上一页/下一页按钮和页码导航。'
         '就像在淘宝浏览商品一样自然。'),
        ('二次筛选', '搜索结果页面保留了筛选条件入口，用户可以在搜索结果内部继续细化筛选。'),
        ('搜索历史管理', '系统自动保存最近 8 条搜索记录在浏览器本地存储中，用户可以点击历史快速搜索，'
         '也可以逐条删除不想保留的历史记录。'),
    ]
    for title, desc in points:
        add_key_value(doc, title, desc)

    add_body(doc, '\n【技术实现要点（给懂技术的人参考）】')
    add_code_block(doc, '• 使用 React 函数式组件 + Hooks（useState, useEffect）管理搜索状态')
    add_code_block(doc, '• 搜索调用 artifactService.searchArtifacts() → 后端 /search 接口')
    add_code_block(doc, '• 搜索历史持久化到 localStorage，下次打开页面自动加载')
    add_code_block(doc, '• 分页逻辑支持"省略号模式"（如 1 ··· 5 6 7 ··· 20）')

    # --- 1.3 ---
    add_styled_heading(doc, '1.3  高级查询（组合条件筛选）', level=2)

    add_body(doc,
        '📄 对应文件：src/pages/AdvancedSearchPage.tsx（约 298 行代码）\n\n'
        '【功能通俗描述】\n'
        '高级查询就像一个"多条件筛选面板"。用户可以通过多个下拉框和输入框，精确限定查询范围。'
        '比如："找出大都会博物馆收藏的、唐代的、青铜材质的文物"。设计参考了克利夫兰博物馆的高级搜索功能。\n\n'
        '【支持的筛选维度】'
    )

    add_table_with_style(doc,
        ['筛选维度', '控件类型', '说明', '示例'],
        [
            ['关键词', '文本输入框', '按文物名称模糊匹配', '输入"菩萨"搜索所有菩萨像'],
            ['文物类型', '下拉选择框', '选择具体文物类别', '雕塑、绘画、陶瓷、青铜器等'],
            ['收藏博物馆', '下拉选择框', '选择收藏机构', '大都会博物馆、大英博物馆等'],
            ['年代/时期', '文本输入 + 年份范围', '支持朝代名称或精确年份', '"唐代" 或 "618—907"'],
            ['年份范围', '两个数字输入框', '精确限定起止年份', '起始 618、结束 907（含公元前负数）'],
        ],
        col_widths=[3, 3, 5, 5]
    )

    add_body(doc, '【核心功能点】')
    points2 = [
        ('左侧筛选面板 + 右侧结果', '页面采用经典的"左右分栏"布局。左侧是固定的筛选条件面板，'
         '右侧展示查询结果。筛选条件变化时结果实时刷新。'),
        ('下拉选项动态加载', '文物类型、博物馆等下拉框的选项不是写死的，而是在页面打开时从后端接口'
         '动态获取，保证选项与数据库实际数据一致，括号里还显示每类有多少件。'),
        ('一键重置', '点击重置按钮可以清空所有筛选条件，回到初始状态。'),
        ('结果导出', '筛选结果同样支持导出为 CSV 或 JSON 文件（详见 1.5 节）。'),
    ]
    for title, desc in points2:
        add_key_value(doc, title, desc)

    # --- 1.4 ---
    add_styled_heading(doc, '1.4  知识图谱查询（关系探索 + 自然语言）', level=2)

    add_body(doc,
        '📄 对应文件：src/pages/KnowledgeGraph.tsx（约 615 行代码，是本模块最复杂的页面）\n\n'
        '【功能通俗描述】\n'
        '知识图谱查询分为两个子标签页：\n'
        '① 图谱浏览：输入一件文物的名字，系统会以"节点 + 连线"的图形方式展示它和其他实体的关系'
        '（比如"罗塞塔石碑 → 收藏于 → 大英博物馆"、"罗塞塔石碑 → 属于 → 古埃及"）。\n'
        '② 自然语言查询：用户可以用日常语言提问，比如输入"唐朝的所有瓷器"，系统会自动解析为查询条件，'
        '找出匹配的文物。'
    )

    add_body(doc, '【图谱浏览模式详细说明】')
    points3 = [
        ('搜索入口', '在搜索框中输入文物名称，系统从后端检索匹配的文物列表，点击选择一件文物即可加载其知识图谱。'),
        ('力导向图展示', '使用 ECharts 图形库绘制"力导向图"。每个节点代表一个实体'
         '（文物 🟦、博物馆 🟨、朝代 🟥、艺术家 🩵、地点 🟩、材质 🟧），'
         '连线代表它们之间的关系（如"收藏于"、"属于"、"创作于"），节点可以被拖拽、图谱可以缩放。'),
        ('节点详情面板', '点击任意节点，右侧面板会显示该节点的详细信息。文物节点显示图片、类型、尺寸、编号；'
         '博物馆节点显示所在地；朝代节点显示时期信息。'),
        ('实体类型图例', '图谱上方自动生成图例，用不同颜色区分节点类型，并标注每类有多少个节点。'),
        ('图谱导出', '知识图谱数据同样支持导出为 CSV 或 JSON，方便学术研究使用。'),
    ]
    for title, desc in points3:
        add_key_value(doc, title, desc)

    add_body(doc, '\n【自然语言查询模式详细说明】')
    points4 = [
        ('输入解析', '用户输入中文自然语言（如"大英博物馆的明代瓷器"），系统通过关键词匹配算法自动识别：'
         '朝代（明 → Ming dynasty）、类型（瓷器 → 陶瓷）、博物馆（大英博物馆 → British Museum）。'),
        ('条件组合', '识别出的条件自动组合为一个高级查询请求，发送给后端获取结果。'),
        ('快捷示例', '页面预设了 5 个查询示例按钮（"唐朝的陶瓷"、"大英博物馆的绘画"等），'
         '一键点击即可体验自然语言查询。'),
        ('结果展示与导出', '查询结果以卡片网格展示，同样支持分页和 CSV/JSON 导出。'),
    ]
    for title, desc in points4:
        add_key_value(doc, title, desc)

    add_body(doc, '\n【自然语言解析的核心逻辑】')
    add_body(doc,
        '系统内置了两张"词典映射表"：\n'
        '• 朝代词典：收录"唐朝 → Tang"、"宋朝 → Song"等 9 个朝代的 27 种中文说法\n'
        '• 类型词典：收录"瓷器 → 陶瓷"、"绘画 → 绘画"等 9 个类别的 18 种中文说法\n'
        '• 博物馆词典：收录"大都会 → Metropolitan Museum of Art"等 7 个博物馆的中文简称\n\n'
        '用户输入后，系统依次在这三张词典中查找匹配项，提取出结构化的查询条件，再调用高级查询接口。'
    )

    # --- 1.5 ---
    add_styled_heading(doc, '1.5  查询结果导出（CSV / JSON）', level=2)

    add_body(doc,
        '【功能通俗描述】\n'
        '在简单查询、高级查询、知识图谱查询三个页面中，查询结果区域右上角都有"导出 CSV"和'
        '"导出 JSON"两个按钮。点击后浏览器会自动下载一个文件，用户可以用 Excel 打开 CSV 文件，'
        '或用文本编辑器打开 JSON 文件。\n\n'
        '【两种导出格式对比】'
    )

    add_table_with_style(doc,
        ['特性', 'CSV 格式', 'JSON 格式'],
        [
            ['文件类型', '表格文件（可用 Excel 打开）', '结构化数据文件（可用记事本/VS Code 打开）'],
            ['适合人群', '普通用户，需要表格处理', '技术人员/研究人员，需要程序处理'],
            ['内容字段', '名称、年代、地区、博物馆、描述（5列）', '文物对象的全部字段（完整 JSON）'],
            ['中文兼容', '自动添加 UTF-8 BOM 头，Excel 打开不乱码', 'UTF-8 编码，含中文无问题'],
            ['技术实现', '前端拼接 CSV 字符串，创建 Blob 下载', 'JSON.stringify + Blob 下载'],
        ],
        col_widths=[3, 6, 6]
    )

    add_body(doc, '\n【导出功能的技术实现（通俗版）】')
    add_body(doc,
        '导出功能完全在浏览器前端完成，不需要后端参与。流程如下：\n'
        '1. 把当前查询结果数据按照目标格式（CSV 或 JSON）组织成文本字符串\n'
        '2. 用这个字符串创建一个"内存中的文件"（Blob 对象）\n'
        '3. 在页面上临时创建一个隐藏的下载链接，模拟点击它\n'
        '4. 浏览器触发下载，文件保存到用户的下载文件夹\n'
        '5. 清理临时对象，释放内存\n\n'
        '整个过程对用户来说就是"点击按钮 → 文件下载完成"。'
    )

    # --- 1.6 ---
    add_styled_heading(doc, '1.6  我的模块涉及的核心文件一览', level=2)

    add_table_with_style(doc,
        ['文件名', '类型', '行数', '功能说明'],
        [
            ['src/pages/SearchPage.tsx', '页面组件', '~300', '简单搜索页面：关键词输入、搜索结果展示、分页、搜索历史管理'],
            ['src/pages/AdvancedSearchPage.tsx', '页面组件', '~298', '高级查询页面：多维度组合筛选、下拉框动态加载、结果展示'],
            ['src/pages/KnowledgeGraph.tsx', '页面组件', '~615', '知识图谱+自然语言查询：力导向图、节点详情、NLQ解析、图谱导出'],
            ['src/pages/BrowsePage.tsx', '页面组件', '~300', '文物浏览页：侧边筛选面板 + 卡片/列表视图切换 + 排序'],
            ['src/services/artifactService.ts', '服务层', '~880', 'API 服务统一入口：封装所有后端接口调用、数据格式转换'],
            ['src/store/artifactStore.ts', '状态管理', '~260', '全局状态仓库：管理文物列表、筛选条件、对比列表等共享状态'],
            ['src/types/artifact.ts', '类型定义', '~57', '定义文物对象、筛选参数、分页响应等数据结构'],
            ['src/types/filter.ts', '类型定义', '~12', '定义筛选状态、视图模式、排序选项等类型'],
            ['src/components/artifacts/FilterPanel.tsx', '筛选组件', '~224', '可复用的多维筛选面板（文物类型、年代、博物馆）'],
            ['src/components/artifacts/ArtifactCard.tsx', '卡片组件', '~186', '文物展示卡片（卡片视图 + 紧凑视图双模式）'],
            ['src/components/artifacts/ArtifactListItem.tsx', '列表组件', '~', '文物展示列表项（列表视图使用）'],
            ['src/components/artifacts/SortControl.tsx', '排序组件', '~', '排序控制（名称/年代/更新时间 + 升降序）'],
        ],
        col_widths=[5.5, 2, 1.3, 7.5]
    )

    doc.add_page_break()

    # ======================= 第二部分 =======================
    add_styled_heading(doc, '第二部分：整个项目的文件架构介绍', level=1)

    # --- 2.1 ---
    add_styled_heading(doc, '2.1  项目技术栈总览', level=2)

    add_body(doc,
        '在介绍文件结构之前，先用通俗语言解释一下项目使用了哪些"建筑材料"：'
    )

    add_table_with_style(doc,
        ['技术名称', '通俗类比', '在项目中的作用'],
        [
            ['React 18', '乐高积木的"底座"', '整个页面的骨架框架，所有界面由它组装而成'],
            ['TypeScript', '给代码加上"标注"', '在写代码时提前标注每个数据的类型，减少出错'],
            ['Vite', '超快的"打包机"', '把代码转换成浏览器能运行的文件，开发时秒级刷新'],
            ['TailwindCSS', '一套"装修模板"', '用预设的样式类名快速美化界面，不需要手写 CSS'],
            ['React Router', '页面的"导航系统"', '控制 URL 与页面之间的跳转关系（如 /search → 搜索页）'],
            ['Zustand', '全局"共享白板"', '多个页面之间共享数据（如筛选条件、对比列表）'],
            ['ECharts', '专业"画图工具"', '绘制饼图、柱状图、力导向图等可视化图表'],
            ['Leaflet', '电子"地图底图"', '展示世界地图和博物馆地理标记'],
            ['Axios / Fetch', '送信的"快递员"', '前端页面与后端服务器之间的数据通信'],
        ],
        col_widths=[3.5, 3.5, 8.5]
    )

    # --- 2.2 ---
    add_styled_heading(doc, '2.2  项目文件结构（逐目录介绍）', level=2)

    add_body(doc, '整个项目的源代码都在 src/ 目录下，共分为 7 个子目录。下面逐层介绍每个文件和目录的作用：')

    # 根目录文件
    add_styled_heading(doc, '项目根目录文件', level=3)
    add_table_with_style(doc,
        ['文件', '作用（通俗版）'],
        [
            ['package.json', '项目的"身份证"——记录了项目名称、版本、依赖了哪些第三方库'],
            ['vite.config.ts', '构建工具的"设置面板"——配置了开发服务器端口(5173)和API代理地址'],
            ['tsconfig.json', 'TypeScript 的"规则手册"——规定了代码类型检查的严格程度'],
            ['tailwind.config.js', '样式系统的"调色板"——定义了自定义颜色（如博物馆金色 #c9a961）'],
            ['index.html', '浏览器加载的"入口页面"——整个应用从这里启动'],
            ['README.md', '项目"说明书"——新人看这个就能了解项目全貌、如何运行'],
            ['CHECKLIST.md', '功能"验收清单"——检查每个功能是否完成'],
        ],
        col_widths=[4.5, 11]
    )

    # src/pages/
    add_styled_heading(doc, 'src/pages/ — 页面目录', level=3)
    add_body(doc, '每个文件对应浏览器中的一个页面（URL 路径）。这是用户直接看到的界面。')
    add_table_with_style(doc,
        ['页面文件', 'URL 路径', '功能简述'],
        [
            ['HomePage.tsx', '/ (首页)', '展示精选文物、特色功能介绍的欢迎页面'],
            ['BrowsePage.tsx', '/browse', '文物浏览页：侧边筛选 + 卡片/列表双视图'],
            ['SearchPage.tsx', '/search', '【我的模块】简单搜索页：关键词全文检索'],
            ['AdvancedSearchPage.tsx', '/advanced-search', '【我的模块】高级查询页：多维度组合筛选'],
            ['KnowledgeGraph.tsx', '/knowledge-graph', '【我的模块】知识图谱：力导向图 + 自然语言查询'],
            ['DetailPage.tsx', '/artifact/:id', '文物详情页：大图展示 + 知识三元组 + 相关推荐'],
            ['ComparePage.tsx', '/compare', '文物对比页：2-3 件文物属性横向对比表格'],
            ['Statistics.tsx', '/statistics', '统计分析看板：饼图、柱状图、总览卡片'],
            ['Timeline.tsx', '/timeline', '文物时间轴：按朝代的时间线可视化'],
            ['Map.tsx', '/map', '地理分布图：Leaflet 世界地图 + 博物馆标记'],
            ['LoginPage.tsx', '/login', '用户登录页'],
            ['RegisterPage.tsx', '/register', '用户注册页'],
            ['ProfilePage.tsx', '/profile', '个人中心页'],
            ['CollectionsPage.tsx', '/collections', '我的收藏页'],
            ['HistoryPage.tsx', '/history', '浏览记录页'],
        ],
        col_widths=[5.5, 3.2, 7]
    )

    # src/components/
    add_styled_heading(doc, 'src/components/ — 可复用组件目录', level=3)
    add_body(doc, '组件就像"预制件"，可以在多个页面中重复使用，避免重复造轮子。')

    add_table_with_style(doc,
        ['子目录', '包含文件', '作用'],
        [
            ['components/ui/', 'button, card, badge, input, skeleton, empty-state, error-state, loading-skeleton',
             '基础 UI 组件库——按钮、卡片、标签、输入框、骨架屏加载动画、空状态/错误状态提示'],
            ['components/layout/', 'Header, Footer, PageLayout',
             '页面布局组件——顶部导航栏（含搜索入口、用户菜单）、底部页脚、页面容器框架'],
            ['components/artifacts/', 'ArtifactCard, ArtifactListItem, FilterPanel, SortControl',
             '文物业务组件——文物卡片（双模式）、列表项、多维筛选面板、排序控件'],
        ],
        col_widths=[4, 5.5, 6.5]
    )

    # src/services/
    add_styled_heading(doc, 'src/services/ — API 服务层', level=3)
    add_body(doc, '服务层是前端与后端之间的"翻译官"，负责发送请求和转换数据格式。')
    add_key_value(doc, 'artifactService.ts（~880行）',
        '核心服务文件，封装了所有后端 API 调用，包括：\n'
        '• 文物列表查询（getArtifacts）\n'
        '• 文物详情查询（getArtifactById）\n'
        '• 简单搜索（searchArtifacts）→ 对应 /search 接口\n'
        '• 高级搜索（advancedSearch）→ 对应 /search/advanced 接口\n'
        '• 知识图谱邻居查询（getGraphNeighbors）→ 对应 /graph/neighbors/{id} 接口\n'
        '• 筛选选项获取（getFilterOptions）→ 对应 /filters 接口\n'
        '• 相关推荐（getRelatedArtifacts）\n'
        '• 对比查询（compareArtifacts）\n'
        '• 统计数据（getStatistics, getStatsDistribution）\n'
        '• 导出功能（exportSearch）\n'
        '• 还包含数据映射函数，将后端返回的英文数据字段自动翻译为中文前端字段')
    add_key_value(doc, 'visualizationService.ts', '为可视化页面（统计、图谱、时间轴、地图）提供数据处理服务')

    # src/store/
    add_styled_heading(doc, 'src/store/ — 全局状态管理', level=3)
    add_body(doc, '可以把 Store 理解为一个"全局共享白板"。多个页面需要访问同一份数据时，'
               '就把数据写在这个白板上，任何页面都能读写。')
    add_key_value(doc, 'artifactStore.ts（~260行）',
        '管理文物列表、当前文物、筛选条件、视图模式、排序方式、分页状态、对比列表等。'
        '当筛选条件变化时，自动重新请求数据。')
    add_key_value(doc, 'userStore.ts', '管理用户登录状态、个人信息、收藏和浏览记录。')

    # src/types/
    add_styled_heading(doc, 'src/types/ — 类型定义', level=3)
    add_body(doc, 'TypeScript 的类型定义文件，就像"数据结构的说明书"——提前约定好每个数据对象有哪些字段、'
               '每个字段是什么类型（文字/数字/数组），这样写代码时编辑器能自动提示和检查错误。')
    add_key_value(doc, 'artifact.ts', '定义 Artifact（文物）、FilterParams（筛选参数）、PaginatedResponse（分页结果）等核心类型')
    add_key_value(doc, 'filter.ts', '定义筛选状态、视图模式、排序选项等 UI 相关类型')
    add_key_value(doc, 'user.ts', '定义用户信息、登录请求等用户相关类型')

    # src/mock/
    add_styled_heading(doc, 'src/mock/ — Mock 模拟数据', level=3)
    add_body(doc, '在后端接口还没开发好时，前端需要假数据来测试页面效果，Mock 就是干这个的。'
               '目前项目已对接真实后端 API，Mock 数据作为备用。')
    add_key_value(doc, 'mock/data/ 目录', '包含 artifacts.ts（12条真实文物数据）、categories.ts、'
        'materials.ts、museums.ts、regions.ts、users.ts')
    add_key_value(doc, 'mock/handlers.ts', '模拟后端 API 的请求处理逻辑')

    # 其他文件和目录
    add_styled_heading(doc, '其他文件与目录', level=3)
    add_table_with_style(doc,
        ['文件/目录', '作用'],
        [
            ['src/App.tsx', 'React 应用的"根组件"，整个页面从这里开始渲染'],
            ['src/main.tsx', '浏览器端的"启动脚本"，把 App 挂载到 index.html 的 <div> 上'],
            ['src/router.tsx', '"路由表"——定义 URL 路径和页面组件的对应关系（如 /search → SearchPage）'],
            ['src/index.css', '全局样式文件，引入 TailwindCSS 并定义自定义主题颜色'],
            ['src/lib/utils.ts', '工具函数（如 cn() 用于合并 CSS 类名）'],
            ['src/utils/imageUtils.ts', '图片处理工具（如图片加载失败时的备用图降级链）'],
            ['docs/', '项目文档目录（技术设计文档等）'],
            ['scripts/', '辅助脚本目录'],
            ['node_modules/', '第三方依赖库的安装目录（npm install 自动生成）'],
        ],
        col_widths=[4.5, 11]
    )

    # --- 2.3 ---
    add_styled_heading(doc, '2.3  页面路由一览', level=2)

    add_body(doc, '下图展示了 URL 路径（浏览器地址栏）与页面组件的对应关系：')

    add_table_with_style(doc,
        ['URL 路径', '页面名称', '所属模块'],
        [
            ['/', '首页', '公共展示'],
            ['/browse', '文物浏览', '数据浏览（含筛选/排序/视图切换）'],
            ['/search', '简单搜索', '【我的模块】数据查询'],
            ['/advanced-search', '高级查询', '【我的模块】数据查询'],
            ['/knowledge-graph', '知识图谱', '【我的模块】数据查询'],
            ['/artifact/:id', '文物详情', '数据浏览'],
            ['/compare', '文物对比', '数据浏览'],
            ['/statistics', '统计分析', '数据可视化'],
            ['/timeline', '时间轴', '数据可视化'],
            ['/map', '地理分布', '数据可视化'],
            ['/login', '登录', '用户管理'],
            ['/register', '注册', '用户管理'],
            ['/profile', '个人中心', '用户管理'],
            ['/collections', '我的收藏', '用户管理'],
            ['/history', '浏览记录', '用户管理'],
        ],
        col_widths=[4, 3, 5.5]
    )

    # --- 2.4 ---
    add_styled_heading(doc, '2.4  数据如何流转（通俗版）', level=2)

    add_body(doc,
        '为了让非技术人员也能理解整个系统的工作流程，这里用一个用户搜索"唐代瓷器"的完整过程来说明：'
    )

    add_body(doc, '【场景：用户在高级查询页面搜索"唐代的瓷器"】')

    steps = [
        ('1. 用户操作', '用户在高级查询页面（AdvancedSearchPage.tsx）的"文物类型"下拉框中选择"陶瓷"，'
         '在"年代/时期"输入框中输入"唐代"，点击"搜索"按钮。'),
        ('2. 页面响应', 'AdvancedSearchPage 组件收集用户输入的筛选条件，组装成一个参数对象：'
         '{ type: "陶瓷", period: "Tang" }，然后调用 artifactService.advancedSearch()。'),
        ('3. 服务层处理', 'artifactService（服务层）将这个参数对象转换后，通过 HTTP 请求发送给后端服务器'
         '的 /search/advanced 接口。同时设置 10 秒超时保护。'),
        ('4. 后端处理', '后端服务器收到请求，在 5441 条文物数据中进行筛选匹配，'
         '将符合条件的文物记录打包成 JSON 格式返回。'),
        ('5. 数据映射', '服务层收到后端返回的原始数据后，调用 mapListItemToArtifact() 函数，'
         '将英文字段名转换为中文前端所需的格式（如 museum.name → museum 字段）。'),
        ('6. 状态更新', '转换后的数据返回给页面组件，页面更新 results 状态，React 自动重新渲染界面。'),
        ('7. 界面展示', '用户看到搜索框右侧出现了 4 列文物卡片网格，每张卡片显示文物图片、名称、年代、博物馆等信息。'
         '右上角显示"导出 CSV"和"导出 JSON"按钮。'),
    ]
    for title, desc in steps:
        add_key_value(doc, title, desc)

    add_body(doc, '\n【数据流转示意图（文字版）】')
    add_code_block(doc, '  用户界面 (Pages)')
    add_code_block(doc, '      ↓ 调用')
    add_code_block(doc, '  API 服务层 (artifactService.ts) ── 数据格式转换 → 返回前端格式')
    add_code_block(doc, '      ↓ HTTP 请求')
    add_code_block(doc, '  后端服务器 (https://se-cs2305.yazs.top)')
    add_code_block(doc, '      ↓ 查询')
    add_code_block(doc, '  数据库 (5441 条文物 + 6 个博物馆)')

    # --- 2.5 ---
    add_styled_heading(doc, '2.5  功能模块总览', level=2)

    add_body(doc, '整个项目从功能上可以划分为以下模块：')

    add_table_with_style(doc,
        ['模块', '包含页面', '说明'],
        [
            ['(1) 数据浏览', '首页、文物浏览、文物详情、文物对比',
             '文物的展示、浏览、详情查看和多文物对比功能'],
            ['(2) 数据查询 ★', '简单搜索、高级查询、知识图谱查询',
             '【组长负责】多层次的文物检索能力，从简单关键词到自然语言查询'],
            ['(3) 数据可视化', '统计分析、时间轴、地理分布',
             '将文物数据以图表、时间线、地图等形式直观呈现'],
            ['(4) 用户管理', '登录、注册、个人中心、收藏、浏览记录',
             '用户账户管理和个性化功能'],
        ],
        col_widths=[3, 4.5, 7.5]
    )

    add_body(doc,
        '\n★ 标注的是组长本人负责的模块。本模块是整个平台用户使用频率最高、最核心的入口功能。'
    )

    # ─── 结尾 ───
    doc.add_page_break()
    add_styled_heading(doc, '附注', level=1)
    add_body(doc, '• 本文档专为汇报展示编写，力求以通俗易懂的语言解释技术实现。')
    add_body(doc, '• 所有代码行数均为约数，实际可能因后续更新而略有差异。')
    add_body(doc, f'• 文档生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
    add_body(doc, '• 项目后端 API 已部署于：https://se-cs2305.yazs.top（5441条真实文物数据 + 6个博物馆）')
    add_body(doc, '• 前端运行地址（开发环境）：http://localhost:5173')

    # ─── 保存 ───
    output_path = r'f:\软件工程大作业\knowledge-service-web-main\数据查询模块_汇报展示文档.docx'
    doc.save(output_path)
    print(f'文档已生成：{output_path}')
    return output_path


if __name__ == '__main__':
    build_document()

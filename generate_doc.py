"""
生成用户个人信息管理模块功能说明文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# ===== 样式设置 =====
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

# 标题样式
for i in range(1, 5):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)

# ===== 封面 =====
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('海外文物知识服务系统\n用户个人信息管理模块\n功能说明文档')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f'文档生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}\n').font.size = Pt(12)
info.add_run('所属系统：海外文物知识服务子系统\n').font.size = Pt(12)
info.add_run('技术栈：React + TypeScript + Vite + Zustand + TailwindCSS\n').font.size = Pt(12)

doc.add_page_break()

# ===== 目录 =====
doc.add_heading('目录', level=1)
toc_items = [
    '1. 模块概述',
    '2. 基础功能：用户注册与登录',
    '   2.1 用户注册',
    '   2.2 用户登录',
    '   2.3 个人资料管理',
    '3. 浏览记录（选做）',
    '4. 收藏功能（选做）',
    '   4.1 收藏夹分组管理',
    '   4.2 收藏操作',
    '5. 评论与互动（选做）',
    '   5.1 发表评论',
    '   5.2 回复评论',
    '   5.3 点赞功能',
    '6. 个性化推荐（选做）——重点详解',
    '   6.1 推荐算法原理',
    '   6.2 数据采集：用户兴趣画像构建',
    '   6.3 匹配策略：关键词模糊匹配',
    '   6.4 冷启动处理',
    '   6.5 完整推荐流程',
    '   6.6 相关文物推荐',
    '7. 代码文件结构',
    '8. 技术栈与架构说明',
    '9. 数据存储方案',
    '10. 路由配置',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ===== 1. 模块概述 =====
doc.add_heading('1. 模块概述', level=1)
doc.add_paragraph(
    '用户个人信息管理模块是海外文物知识服务系统的核心子系统之一，负责处理所有与用户相关的功能。'
    '该模块采用前后端分离架构（当前为纯前端实现，使用Mock API模拟后端），'
    '实现了从用户注册登录到个性化内容推荐的完整用户生命周期管理。'
)
doc.add_paragraph('本模块包含以下五大功能板块：')

features = [
    ('基础功能（必做）', '用户注册、登录、个人资料管理（用户名、密码、头像、昵称、简介）'),
    ('浏览记录（选做）', '自动记录用户浏览的文物，支持去重和时间排序，支持一键清空'),
    ('收藏功能（选做）', '支持用户收藏文物，收藏夹支持分组管理（创建、编辑、删除分组，移动收藏项）'),
    ('评论与互动（选做）', '支持用户对文物发表评论，支持点赞与回复'),
    ('个性化推荐（选做）', '基于用户浏览与收藏记录，在首页推送可能感兴趣的文物'),
]
for name, desc in features:
    p = doc.add_paragraph()
    run = p.add_run(f'• {name}：')
    run.font.bold = True
    p.add_run(desc)

doc.add_paragraph()

# ===== 2. 基础功能：用户注册与登录 =====
doc.add_heading('2. 基础功能：用户注册与登录', level=1)

doc.add_heading('2.1 用户注册', level=2)
doc.add_paragraph(
    '注册页面（RegisterPage.tsx）提供完整的用户注册表单，支持以下字段：'
)
fields = [
    '用户名（必填）：3-20个字符',
    '昵称（选填）：用于展示的名称',
    '邮箱（必填）：需符合邮箱格式验证',
    '密码（必填）：至少6位',
    '确认密码（必填）：需与密码一致',
]
for f in fields:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('注册流程的关键代码逻辑：', style='Intense Quote')

code_text = '''// RegisterPage.tsx - handleSubmit 函数
const handleSubmit = async (e: React.FormEvent): Promise<void> => {
  e.preventDefault();
  // 1. 输入验证：检查用户名长度、邮箱格式、密码长度、两次密码一致性
  if (!username.trim()) { setLocalError('请输入用户名'); return; }
  if (username.trim().length < 3) { setLocalError('用户名至少需要3个字符'); return; }
  // ... 更多验证

  // 2. 调用 Zustand Store 的 register 方法
  const success = await register({
    username: username.trim(),
    email: email.trim(),
    password,
    nickname: nickname.trim() || undefined,
  });

  // 3. 注册成功后跳转首页
  if (success) { navigate('/'); }
};'''

p = doc.add_paragraph()
run = p.add_run(code_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph(
    '注册请求发送到 Mock API（handlers.ts 中的 userApi.register），Mock API 会：\n'
    '1) 检查用户名和邮箱是否已存在\n'
    '2) 对密码进行哈希处理（simpleHash 函数）\n'
    '3) 将新用户存入内存中的 usersStore（Map 数据结构）\n'
    '4) 自动为新用户创建默认收藏分组\n'
    '5) 生成认证 token 并返回'
)

doc.add_heading('2.2 用户登录', level=2)
doc.add_paragraph(
    '登录页面（LoginPage.tsx）支持使用用户名或邮箱登录，提供密码显示/隐藏切换功能。'
    '登录成功后，认证信息（token 和用户数据）会持久化到 localStorage，'
    '使得用户刷新页面后仍然保持登录状态。'
)

doc.add_paragraph('登录状态管理的关键实现：')
code_text2 = '''// userStore.ts - 状态初始化（从 localStorage 恢复登录状态）
const loadPersistedState = <T>(key: string, defaultValue: T): T => {
  try {
    const saved = localStorage.getItem(key);
    if (saved) { return JSON.parse(saved) as T; }
  } catch { /* ignore */ }
  return defaultValue;
};

// 初始化时自动恢复登录状态
currentUser: loadPersistedState<UserProfile | null>(CURRENT_USER_KEY, null),
token: localStorage.getItem(AUTH_TOKEN_KEY),
isAuthenticated: !!localStorage.getItem(AUTH_TOKEN_KEY),'''
p = doc.add_paragraph()
run = p.add_run(code_text2)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('2.3 个人资料管理', level=2)
doc.add_paragraph(
    '个人资料页面（ProfilePage.tsx）提供以下功能：\n'
    '• 头像管理：点击头像按钮循环切换预设头像\n'
    '• 昵称与简介编辑：内联编辑表单，支持保存和取消\n'
    '• 密码修改：需要输入旧密码和新密码\n'
    '• 统计展示：显示收藏数量和浏览记录数量\n'
    '• 快捷入口：链接到收藏、浏览记录、文物浏览、文物对比等页面\n'
    '• 退出登录：清除所有本地存储的认证和用户数据'
)

doc.add_page_break()

# ===== 3. 浏览记录 =====
doc.add_heading('3. 浏览记录（选做）', level=1)
doc.add_paragraph(
    '浏览记录功能自动追踪用户查看过的文物，提供便捷的历史回溯能力。'
)

doc.add_heading('3.1 自动记录机制', level=2)
doc.add_paragraph(
    '当用户在文物详情页（DetailPage.tsx）查看文物时，系统自动调用 addBrowseHistory 方法记录浏览行为：'
)

code_text3 = '''// DetailPage.tsx - useEffect 中的数据加载逻辑
useEffect(() => {
  const loadData = async (): Promise<void> => {
    const [artifactData, recommendationsData] = await Promise.all([
      artifactService.getArtifactById(id),
      artifactService.getRelatedArtifacts(id, 4)
    ]);
    setArtifact(artifactData);
    // 记录浏览历史（仅在用户已登录时）
    if (artifactData && isAuthenticated) {
      addBrowseHistory(
        artifactData.id,
        artifactData.name,
        artifactData.images[0] || '',
        artifactData.category,
        artifactData.era
      );
    }
  };
  loadData();
}, [id]);'''
p = doc.add_paragraph()
run = p.add_run(code_text3)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('3.2 浏览记录管理', level=2)
doc.add_paragraph(
    'addBrowseHistory 方法实现了以下逻辑：\n'
    '1) 去重：如果同一文物已存在于浏览记录中，先移除旧的记录\n'
    '2) 插入：将新记录插入到列表最前面\n'
    '3) 限长：保持最多50条记录，超出则删除最旧的\n'
    '4) 持久化：将记录保存到 localStorage\n\n'
    '浏览记录页面（HistoryPage.tsx）提供：\n'
    '• 按时间倒序排列的浏览列表\n'
    '• 人性化的时间显示（"刚刚"、"5分钟前"、"3天前"等）\n'
    '• 一键清空所有记录（需确认）\n'
    '• 每条记录可点击跳转到对应文物详情页'
)

doc.add_page_break()

# ===== 4. 收藏功能 =====
doc.add_heading('4. 收藏功能（选做）', level=1)
doc.add_paragraph(
    '收藏功能允许用户将感兴趣的文物保存到收藏夹中，并支持分组管理。'
)

doc.add_heading('4.1 收藏夹分组管理', level=2)
doc.add_paragraph(
    '用户可以为收藏的文物创建多个收藏夹（如"中国陶瓷"、"欧洲绘画"等），方便分类管理。'
    '系统支持以下分组操作：\n'
    '• 创建收藏夹：设置名称和描述\n'
    '• 编辑收藏夹：修改名称和描述\n'
    '• 删除收藏夹：同时删除该分组下的所有收藏项\n'
    '• 默认收藏夹：新用户自动创建"默认收藏夹"'
)

doc.add_heading('4.2 收藏操作', level=2)
doc.add_paragraph(
    '在文物详情页，用户可以点击"收藏"按钮（心形图标）来收藏/取消收藏文物：\n'
    '• 未收藏状态：灰色空心心形，点击后变为红色实心心形\n'
    '• 已收藏状态：红色实心心形，点击后取消收藏\n'
    '• 未登录用户点击收藏时会跳转到登录页\n\n'
    '在收藏管理页面（CollectionsPage.tsx）中，用户可以：\n'
    '• 按分组查看收藏的文物\n'
    '• 将收藏项移动到其他分组\n'
    '• 取消收藏（移除收藏项）\n'
    '• 点击缩略图跳转到文物详情页'
)

doc.add_paragraph()
doc.add_paragraph('收藏功能的核心状态管理（userStore.ts）：', style='Intense Quote')

code_text4 = '''// userStore.ts - toggleCollectArtifact（收藏/取消收藏切换）
toggleCollectArtifact: async (artifactId, artifactName, ...) => {
  const { currentUser, collectionGroups, collectionItems } = get();
  if (!currentUser) return false;

  try {
    // 如果没有收藏夹，通过 API 自动创建默认收藏夹
    let groups = collectionGroups;
    if (groups.length === 0) {
      const defaultGroup = await userApi.createCollectionGroup(
        currentUser.id, '默认收藏夹', '我的默认收藏夹'
      );
      groups = [...groups, defaultGroup];
    }

    // 检查是否已收藏
    const existing = collectionItems.find(item => item.artifactId === artifactId);
    if (existing) {
      await removeFromCollection(existing.id);  // 取消收藏
      return false;
    }

    // 添加到第一个（默认）收藏夹
    await addToCollection(groups[0].id, artifactId, artifactName, ...);
    return true;
  } catch (err) { /* 错误处理 */ }
};'''
p = doc.add_paragraph()
run = p.add_run(code_text4)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph(
    '⚠️ 修复说明：原始代码存在数据不同步问题——toggleCollectArtifact 直接修改 Zustand 状态和 localStorage，'
    '但当用户进入"我的收藏"页面时，fetchCollectionGroups/fetchCollectionItems 从 Mock API 获取数据，'
    '覆盖了之前的状态。修复方案是将所有收藏操作改为调用 Mock API，确保数据源统一。'
)

doc.add_page_break()

# ===== 5. 评论与互动 =====
doc.add_heading('5. 评论与互动（选做）', level=1)

doc.add_heading('5.1 发表评论', level=2)
doc.add_paragraph(
    '在文物详情页底部，登录用户可以通过输入框发表评论。评论支持 Enter 键快捷发送。'
    '未登录用户会看到"登录后即可发表评论"的提示。'
    '评论数据结构包含：评论ID、文物ID、用户信息、内容、点赞数、回复列表等。'
)

doc.add_heading('5.2 回复评论', level=2)
doc.add_paragraph(
    '用户可以点击评论下方的"回复"按钮展开回复区域，对任意评论进行回复。'
    '回复支持嵌套显示，每条评论下可以有多条回复。'
)

doc.add_heading('5.3 点赞功能', level=2)
doc.add_paragraph(
    '用户可以对评论进行点赞/取消点赞操作：\n'
    '• 点击拇指图标切换点赞状态\n'
    '• 使用 likedBy 数组记录点赞用户ID，防止重复点赞\n'
    '• 点赞数实时更新显示\n'
    '• 当前用户的点赞状态用金色高亮显示'
)

doc.add_paragraph()
doc.add_paragraph('评论功能的代码位置：', style='Intense Quote')
doc.add_paragraph('• 添加评论：userStore.ts → addComment() 方法（第464-478行）')
doc.add_paragraph('• 添加回复：userStore.ts → addReply() 方法（第480-511行）')
doc.add_paragraph('• 点赞切换：userStore.ts → likeComment() 方法（第513-547行）')
doc.add_paragraph('• UI 渲染：DetailPage.tsx → 评论区域（第565-741行）')

doc.add_page_break()

# ===== 6. 个性化推荐 =====
doc.add_heading('6. 个性化推荐（选做）——重点详解', level=1)
doc.add_paragraph(
    '个性化推荐是本模块最核心的选做功能，它根据用户的浏览历史和收藏记录，'
    '在首页"猜你喜欢"区域智能推送用户可能感兴趣的文物。以下从算法原理、数据采集、'
    '匹配策略、冷启动处理等方面进行详细解析。'
)

doc.add_heading('6.1 推荐算法原理', level=2)
doc.add_paragraph(
    '本系统采用基于内容的推荐算法（Content-Based Filtering），核心思想是：\n'
    '"如果用户过去喜欢（浏览/收藏）某些文物，那么推荐与这些文物属性相似的其他文物。"\n\n'
    '算法的核心流程分为三个阶段：\n'
    '1) 用户兴趣画像构建：从用户的浏览记录和收藏记录中提取感兴趣的属性关键词\n'
    '2) 候选文物匹配：将兴趣关键词与文物库中的文物属性进行模糊匹配\n'
    '3) 冷启动处理：当用户兴趣数据不足时，使用热门文物补齐推荐列表'
)

doc.add_heading('6.2 数据采集：用户兴趣画像构建', level=2)
doc.add_paragraph(
    '系统从两个数据源采集用户兴趣信息：浏览记录（browseHistory）和收藏记录（collectionItems）。'
    '对每一条记录，提取两类关键词：'
)

doc.add_paragraph('A. 朝代关键词提取', style='List Bullet')
doc.add_paragraph(
    '系统维护了一个预定义的朝代名称列表（DYNASTY_NAMES），包含中英文朝代名：\n'
    "'Qing', 'Ming', 'Tang', 'Song', 'Han', 'Yuan', 'Zhou', 'Shang', 'Qin',\n"
    "'清朝', '明朝', '唐朝', '宋朝', '汉朝', '元朝', '周朝', '商朝', '秦朝'\n\n"
    '遍历用户浏览/收藏的每件文物的 era（年代）字段，检查是否包含上述朝代关键词。'
    '例如，文物的 era 字段为"清朝 (1644–1911)"，则提取出关键词"清朝"。'
)

doc.add_paragraph('B. 类别关键词提取', style='List Bullet')
doc.add_paragraph(
    '提取文物的 category（类别）字段作为关键词。过滤掉"其他"和"未知"等无意义类别。'
    '例如，文物类别为"瓷器"，则提取关键词"瓷器"。'
)

doc.add_paragraph()
doc.add_paragraph('兴趣画像构建的核心代码（HomePage.tsx 第40-80行）：', style='Intense Quote')

code_text5 = '''// HomePage.tsx - 用户兴趣关键词提取
const interestKeywords: string[] = [];
const seenIds = new Set<string>();

const DYNASTY_NAMES = ['Qing', 'Ming', 'Tang', 'Song', 'Han', 'Yuan', 'Zhou',
                       'Shang', 'Qin', '清朝', '明朝', '唐朝', '宋朝', '汉朝',
                       '元朝', '周朝', '商朝', '秦朝'];

// 从浏览记录中提取关键词
browseHistory.forEach((h) => {
  seenIds.add(h.artifactId);  // 记录已浏览的文物ID（避免推荐重复）
  if (h.artifactEra) {
    for (const name of DYNASTY_NAMES) {
      if (h.artifactEra.includes(name) && !interestKeywords.includes(name)) {
        interestKeywords.push(name);
      }
    }
  }
  if (h.artifactCategory && h.artifactCategory !== '其他') {
    if (!interestKeywords.includes(h.artifactCategory)) {
      interestKeywords.push(h.artifactCategory);
    }
  }
});

// 从收藏记录中同样提取关键词（代码结构相同）
collectionItems.forEach((item) => {
  seenIds.add(item.artifactId);
  // ... 同样的关键词提取逻辑
});'''
p = doc.add_paragraph()
run = p.add_run(code_text5)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('6.3 匹配策略：关键词模糊匹配', level=2)
doc.add_paragraph(
    '获取到用户的兴趣关键词列表后，系统从文物库中获取前50件文物作为候选集，'
    '然后进行两轮筛选：'
)

doc.add_paragraph('第一轮：兴趣匹配', style='List Bullet')
doc.add_paragraph(
    '从候选文物中筛选出满足以下条件的文物：\n'
    '• 文物尚未被用户浏览或收藏（不在 seenIds 中）\n'
    '• 文物的 era（年代）字段包含任意一个兴趣关键词，或\n'
    '  文物的 category（类别）字段包含任意一个兴趣关键词\n'
    '• 最多取4件作为推荐结果'
)

doc.add_paragraph()
doc.add_paragraph('匹配代码（HomePage.tsx 第90-98行）：', style='Intense Quote')

code_text6 = '''if (interestKeywords.length > 0) {
  recommendations = allArtifacts
    .filter((artifact) => !seenIds.has(artifact.id))    // 排除已看过的
    .filter((artifact) => {
      return interestKeywords.some(kw =>
        artifact.era.includes(kw) ||                      // 年代匹配
        artifact.category.includes(kw)                     // 类别匹配
      );
    })
    .slice(0, 4);                                         // 取前4条
}'''
p = doc.add_paragraph()
run = p.add_run(code_text6)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph('第二轮：随机补齐（冷启动处理）', style='List Bullet')
doc.add_paragraph(
    '如果第一轮匹配结果不足4条，从剩余候选中随机选取补齐：'
)

code_text7 = '''if (recommendations.length < 4) {
  const remaining = allArtifacts
    .filter(a => !seenIds.has(a.id)
              && !recommendations.find(r => r.id === a.id))
    .sort(() => 0.5 - Math.random())    // 随机打乱
    .slice(0, 4 - recommendations.length);
  recommendations = [...recommendations, ...remaining];
}'''
p = doc.add_paragraph()
run = p.add_run(code_text7)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('6.4 冷启动处理', level=2)
doc.add_paragraph(
    '冷启动（Cold Start）是推荐系统面临的核心挑战之一。当新用户没有任何浏览和收藏记录时，'
    '系统无法提取兴趣关键词。本系统采用以下策略处理冷启动问题：'
)
doc.add_paragraph(
    '1) 新用户冷启动：当用户无浏览/收藏记录时，interestKeywords 为空数组，'
    '第一轮匹配产生0条结果，完全依赖第二轮的随机热门文物补齐。\n'
    '2) 数据稀疏冷启动：当用户只有少量浏览记录，提取的关键词覆盖面有限时，'
    '第一轮可能产生不足4条结果，由第二轮随机补齐。\n'
    '3) 提示文案自适应：当用户有浏览/收藏记录时显示"基于你的浏览和收藏记录，为你推荐这些文物"；'
    '否则显示"为你精选的热门文物推荐"。'
)

doc.add_paragraph()
doc.add_paragraph('冷启动文案切换（HomePage.tsx 第282-284行）：', style='Intense Quote')

code_text8 = '''{isAuthenticated && (browseHistory.length > 0 || collectionItems.length > 0)
  ? '基于你的浏览和收藏记录，为你推荐这些文物'
  : '为你精选的热门文物推荐'}'''
p = doc.add_paragraph()
run = p.add_run(code_text8)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('6.5 完整推荐流程', level=2)
doc.add_paragraph('个性化推荐的完整执行流程如下：')
doc.add_paragraph(
    'Step 1: 组件挂载（HomePage.tsx 第38-118行的 useEffect）\n'
    '        ↓\n'
    'Step 2: 监听 isAuthenticated, browseHistory, collectionItems 变化\n'
    '        ↓\n'
    'Step 3: 遍历 browseHistory 和 collectionItems\n'
    '        ├── 提取朝代关键词（era字段中的朝代名）\n'
    '        ├── 提取类别关键词（category字段）\n'
    '        └── 收集 seenIds（已浏览/收藏的文物ID）\n'
    '        ↓\n'
    'Step 4: 调用 artifactService.getArtifacts({page:1, size:50}) 获取候选集\n'
    '        ↓\n'
    'Step 5: 如果有兴趣关键词 → 模糊匹配筛选（era或category包含关键词）\n'
    '        ↓\n'
    'Step 6: 如果不足4条 → 随机补齐热门文物\n'
    '        ↓\n'
    'Step 7: 设置 recommendedArtifacts 状态 → 渲染"猜你喜欢"区域'
)

doc.add_heading('6.6 相关文物推荐（辅助推荐）', level=2)
doc.add_paragraph(
    '除首页的个性化推荐外，文物详情页底部还提供了"相关文物推荐"功能。'
    '该推荐基于后端API（getRelatedArtifacts），采用加权评分算法：\n'
    '• 同地区（region）：权重 3.0\n'
    '• 同类别（category）：权重 2.5\n'
    '• 同时代（era）：权重 2.5\n'
    '• 同材质（material）：权重 2.0\n'
    '• 同博物馆（museum）：权重 1.5\n'
    '• 共同标签（tags）：每个 1.5\n\n'
    '该推荐与个性化推荐的互补关系：首页推荐侧重"用户长期兴趣"，'
    '详情页推荐侧重"当前文物的相似性"。'
)

doc.add_page_break()

# ===== 7. 代码文件结构 =====
doc.add_heading('7. 代码文件结构', level=1)
doc.add_paragraph('用户个人信息管理模块涉及的核心文件：')

file_structure = [
    ('页面层（Pages）', [
        ('src/pages/LoginPage.tsx', '用户登录页面，支持用户名/邮箱登录、密码显隐切换'),
        ('src/pages/RegisterPage.tsx', '用户注册页面，含完整表单验证逻辑'),
        ('src/pages/ProfilePage.tsx', '个人资料页面，头像、昵称、简介编辑，密码修改'),
        ('src/pages/CollectionsPage.tsx', '收藏管理页面，分组列表+收藏项网格+CRUD操作'),
        ('src/pages/HistoryPage.tsx', '浏览记录页面，时间排序+清空操作'),
        ('src/pages/HomePage.tsx', '首页"猜你喜欢"个性化推荐区域'),
        ('src/pages/DetailPage.tsx', '文物详情页，含收藏按钮+评论区域+浏览记录触发'),
    ]),
    ('状态管理层（Store）', [
        ('src/store/userStore.ts', 'Zustand状态管理核心，管理认证、收藏、评论、浏览记录等所有用户状态'),
    ]),
    ('Mock数据与API层', [
        ('src/mock/handlers.ts', 'Mock API实现，含userApi（登录、注册、收藏CRUD、评论获取）'),
        ('src/mock/data/users.ts', '预设用户数据（3个测试账号）+ 评论数据存储'),
    ]),
    ('类型定义层（Types）', [
        ('src/types/user.ts', 'TypeScript类型定义：User, CollectionGroup, CollectionItem, Comment等'),
    ]),
    ('路由配置', [
        ('src/router.tsx', '所有页面路由定义，用户相关路由包括/login, /register, /profile, /collections, /history'),
    ]),
]

for category, files in file_structure:
    doc.add_heading(category, level=3)
    for file_path, desc in files:
        p = doc.add_paragraph()
        run = p.add_run(f'{file_path}')
        run.font.bold = True
        run.font.size = Pt(10)
        p.add_run(f'\n    {desc}')
        p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ===== 8. 技术栈 =====
doc.add_heading('8. 技术栈与架构说明', level=1)

doc.add_heading('8.1 前端技术栈', level=2)
tech_items = [
    ('React 18.3', 'UI框架，采用函数式组件 + Hooks 模式'),
    ('TypeScript 5.5', '类型安全，所有接口和数据模型均有完整类型定义'),
    ('Vite 5.4', '构建工具，支持HMR热更新，开发体验极佳'),
    ('Zustand 4.5', '轻量级状态管理，用于管理用户认证、收藏、评论等全局状态'),
    ('React Router 6.26', '客户端路由，支持嵌套路由和参数传递'),
    ('TailwindCSS 3.4', '原子化CSS框架，自定义博物馆主题色彩'),
    ('Axios 1.16', 'HTTP客户端（已安装，用于后期对接真实后端）'),
    ('lucide-react', '图标库，提供一致的UI图标系统'),
]
for name, desc in tech_items:
    p = doc.add_paragraph()
    run = p.add_run(f'• {name}：')
    run.font.bold = True
    p.add_run(desc)

doc.add_heading('8.2 架构模式', level=2)
doc.add_paragraph(
    '项目采用分层架构设计：\n\n'
    '┌─────────────────────────────────────┐\n'
    '│        Pages（页面层）                │\n'
    '│  LoginPage, CollectionsPage, etc.   │\n'
    '├─────────────────────────────────────┤\n'
    '│     Components（组件层）              │\n'
    '│   ArtifactCard, FilterPanel, etc.   │\n'
    '├─────────────────────────────────────┤\n'
    '│    Store（状态管理层）Zustand         │\n'
    '│   userStore.ts, artifactStore.ts    │\n'
    '├─────────────────────────────────────┤\n'
    '│    Services（服务层）                 │\n'
    '│  artifactService.ts (真实API)        │\n'
    '│  mock/handlers.ts (Mock API)        │\n'
    '├─────────────────────────────────────┤\n'
    '│    Types（类型定义层）                 │\n'
    '│   user.ts, artifact.ts, filter.ts   │\n'
    '└─────────────────────────────────────┘\n\n'
    '数据流向：Page → Store (action) → API (Mock/Real) → Store (state update) → Page (re-render)'
)

doc.add_heading('8.3 当前运行模式', level=2)
doc.add_paragraph(
    '• 文物数据：通过 Vite 代理转发到远程知识图谱子系统 API（https://se-cs2305.yazs.top）\n'
    '• 用户数据：使用 Mock API（handlers.ts 中的 userApi），数据存储在内存和 localStorage 中\n'
    '• 开发服务器：npm run dev → http://localhost:5173\n'
    '• 如需完全离线运行：修改 artifactService.ts 使用 Mock API 替代真实 API'
)

doc.add_page_break()

# ===== 9. 数据存储方案 =====
doc.add_heading('9. 数据存储方案', level=1)
doc.add_paragraph(
    '当前系统（前端阶段）采用双层存储策略：'
)

table = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
table.autofit = True
headers = ['数据类型', '内存存储（Mock API）', 'localStorage（持久化）']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

data_rows = [
    ('用户数据', 'usersStore (Map)', 'CURRENT_USER_KEY → currentUser'),
    ('认证Token', '—', 'AUTH_TOKEN_KEY → token'),
    ('浏览记录', 'browseHistoryStore (Map)', 'browse_history → browseHistory[]'),
    ('收藏分组', 'collectionGroupsStore (Map)', 'collection_groups → CollectionGroup[]'),
    ('收藏项', 'collectionItemsStore (Map)', 'collection_items → CollectionItem[]'),
    ('评论数据', 'commentsStore (Array)', 'comments → Record<string, Comment[]>'),
]
for i, (data_type, mem, local) in enumerate(data_rows):
    table.rows[i+1].cells[0].text = data_type
    table.rows[i+1].cells[1].text = mem
    table.rows[i+1].cells[2].text = local

doc.add_paragraph()
doc.add_paragraph(
    '⚠️ 重要说明：localStorage 用于页面刷新后的状态恢复，Mock API 的内存存储用于模拟后端数据库。'
    '在修复后的代码中，所有数据操作都先写入 Mock API（内存），然后同步到 Zustand Store 和 localStorage，'
    '确保数据一致性。'
)

doc.add_page_break()

# ===== 10. 路由配置 =====
doc.add_heading('10. 路由配置', level=1)
doc.add_paragraph('用户个人信息管理模块涉及的路由：')

routes = [
    ('/', 'HomePage', '首页（含个性化推荐"猜你喜欢"区域）'),
    ('/login', 'LoginPage', '用户登录页'),
    ('/register', 'RegisterPage', '用户注册页'),
    ('/profile', 'ProfilePage', '个人资料管理页'),
    ('/collections', 'CollectionsPage', '我的收藏页（收藏夹分组管理+收藏项管理）'),
    ('/history', 'HistoryPage', '浏览记录页'),
    ('/artifact/:id', 'DetailPage', '文物详情页（含收藏按钮+评论区）'),
]
for path, component, desc in routes:
    p = doc.add_paragraph()
    run = p.add_run(f'{path}')
    run.font.bold = True
    run.font.name = 'Consolas'
    p.add_run(f'  →  {component}  —— {desc}')

doc.add_paragraph()

# ===== 总结 =====
doc.add_heading('总结', level=1)
doc.add_paragraph(
    '用户个人信息管理模块实现了完整的用户生命周期管理功能，包括：\n\n'
    '1. 基础功能（必做）：用户注册、登录、个人资料管理 —— 全部完成 ✓\n\n'
    '2. 浏览记录（选做）：自动记录、去重、排序、一键清空 —— 全部完成 ✓\n\n'
    '3. 收藏功能（选做）：分组管理、CRUD操作、移动收藏项 —— 全部完成 ✓\n'
    '   （已修复数据不同步的Bug，确保 Mock API 与 localStorage 数据一致）\n\n'
    '4. 评论与互动（选做）：发表评论、回复、点赞/取消点赞 —— 全部完成 ✓\n\n'
    '5. 个性化推荐（选做）：基于内容的推荐算法、关键词提取与模糊匹配、冷启动处理 —— 全部完成 ✓\n\n'
    '该模块采用现代前端技术栈（React 18 + TypeScript + Zustand + TailwindCSS），'
    '代码结构清晰、类型安全、可维护性强，为后续对接真实后端奠定了良好基础。'
)

# ===== 保存 =====
output_path = r'c:\Users\houtengdu\Desktop\大作业重生\knowledge-service-web-main (1)\knowledge-service-web-main\用户个人信息管理模块-功能说明文档.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')
